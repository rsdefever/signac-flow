# Copyright (c) 2018 The Regents of the University of Michigan
# All rights reserved.
# This software is licensed under the BSD 3-Clause License.

"""This is a helper script to generate a .docx file containing all templates for easy review."""
import os
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import git

import signac

import extract_templates as ext

DOC_DIR = os.path.join(
    os.path.dirname(__file__), 'compiled_scripts')



PROJECT_DEFINITION = """
class TestProject(flow.FlowProject):
    ngpu=2
    np=3
    omp_num_threads=4
    nranks=5


@TestProject.operation
def serial_op(job):
    pass


@TestProject.operation
@flow.directives(np=TestProject.np)
def parallel_op(job):
    pass


@TestProject.operation
@flow.directives(nranks=TestProject.nranks)
def mpi_op(job):
    pass


@TestProject.operation
@flow.directives(omp_num_threads=TestProject.omp_num_threads)
def omp_op(job):
    pass


@TestProject.operation
@flow.directives(nranks=TestProject.nranks, omp_num_threads=TestProject.omp_num_threads)
def hybrid_op(job):
    pass


@TestProject.operation
@flow.directives(ngpu=TestProject.ngpu)
def gpu_op(job):
    pass


@TestProject.operation
@flow.directives(ngpu=TestProject.ngpu, nranks=TestProject.nranks)
def mpi_gpu_op(job):
    pass

"""

def add_param_header(document, job):
    h = ['{}={}'.format(key, val) for key, val in job.sp.parameters.items()]
    h = 'Parameters: ' + ', '.join(h)
    document.add_heading(h, level=2)


def process_job(document, job):
    add_param_header(document, job)

    from collections import OrderedDict
    name_map = OrderedDict([
            ('script_serial_op.sh',  'Serial operation'),
            ('script_parallel_op.sh',  'Generic parallel operation'),
            ('script_mpi_op.sh',  'MPI operation'),
            ('script_omp_op.sh',  'OpenMP operation'),
            ('script_hybrid_op.sh',  'MPI-OpenMP Hybrid operation'),
            ('script_gpu_op.sh',  'GPU operation'),
            ('script_mpi_gpu_op.sh',  'MPI-GPU operation'),
            ])

    if job.sp.parameters.get('bundle', False):
        h = 'Bundled MPI and OpenMP jobs'
        for fn in os.listdir(job.workspace()):
            if fn == 'signac_statepoint.json':
                continue
            document.add_heading(h, level=3)
            with open(job.fn(fn)) as fr:
                p = document.add_paragraph(fr.read(), style='Code')
    else:
        for fn, h in name_map.items():
            document.add_heading('Operation: {}'.format(h), level=3)
            if job.isfile(fn):
                with open(job.fn(fn)) as fr:
                    p = document.add_paragraph(fr.read(), style='Code')
            else:
                print('{}: {} is missing.'.format(job, fn))
                p = document.add_paragraph('Script {} not present!'.format(job.fn(fn)))

    document.add_page_break()


def main():
    repo = git.Repo('..')
    commit = repo.head.commit

    if not os.path.exists(DOC_DIR):
        os.makedirs(DOC_DIR)
    project = signac.get_project(ext.PROJECT_DIR)

    environments = project.detect_schema()['environment'][str]
    for env in sorted(environments):
        env_name = env.split('.')[-1]
        print('Generating review document for {}...'.format(env_name))
        document = docx.Document()

        # Add code style
        style = document.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Monaco'
        style.font.size = Pt(8)
        style = document.styles.add_style('CodeChar', WD_STYLE_TYPE.CHARACTER)
        style.font.name = 'Monaco'
        style.font.size = Pt(8)

        document.add_heading(env_name, level=0)
        p = document.add_paragraph("Output at commit ")
        p.add_run('{}'.format(commit), style='CodeChar')
        document.add_heading("FlowProject Definition", level=1)
        p = document.add_paragraph(PROJECT_DEFINITION, style='Code')
        document.add_page_break()

        document.add_heading("Operations without bundling", level=1)
        query = {'environment': env, 'parameters.bundle': {'$exists': False}}
        for job in project.find_jobs(query):
            process_job(document, job)

        document.add_heading("Operations with bundling", level=1)
        query = {'environment': env, 'parameters.bundle': {'$exists': True}}
        for job in project.find_jobs(query):
            process_job(document, job)


        fn = os.path.join(DOC_DIR, "{env}.docx".format(env=env_name))
        document.save(fn)

if __name__ == "__main__":
    main()
